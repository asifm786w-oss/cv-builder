from __future__ import annotations

from io import BytesIO
from pathlib import Path
from datetime import date
from urllib.parse import quote
import asyncio
import logging
import os
import sys

import psycopg2
from jinja2 import Environment, FileSystemLoader, select_autoescape
from docx import Document

from models import CV

logger = logging.getLogger(__name__)

# ============================================================
# DB CHECK
# ============================================================
def verify_postgres_connection() -> bool:
    db_url = os.getenv("DATABASE_URL")

    if not db_url:
        logger.error("[DB CHECK] DATABASE_URL is NOT set")
        return False

    try:
        conn = psycopg2.connect(db_url, connect_timeout=5)
        cur = conn.cursor()
        cur.execute("SELECT 1;")
        result = cur.fetchone()
        cur.close()
        conn.close()

        logger.info(f"[DB CHECK] Postgres connected successfully, result={result}")
        return True

    except Exception as e:
        logger.exception(f"[DB CHECK] Postgres connection FAILED: {e}")
        return False


# ============================================================
# Jinja setup
# ============================================================
BASE_DIR = Path(__file__).resolve().parent
TEMPLATES_DIR = BASE_DIR / "templates"

env = Environment(
    loader=FileSystemLoader(str(TEMPLATES_DIR)),
    autoescape=select_autoescape(["html", "xml"]),
)

def render_cv_html(cv: CV, template_name: str) -> str:
    template = env.get_template(template_name)
    return template.render(cv=cv)


# ============================================================
# PDF (Playwright-only)
# ============================================================
def _prepare_windows_event_loop() -> None:
    if sys.platform.startswith("win"):
        try:
            asyncio.set_event_loop_policy(asyncio.WindowsSelectorEventLoopPolicy())  # type: ignore[attr-defined]
        except Exception as e:
            logger.warning(f"[PDF] Could not set Windows event loop policy: {e}")


# Set a stable browser path (also set as Railway Variable recommended)
os.environ.setdefault("PLAYWRIGHT_BROWSERS_PATH", "/app/.playwright")


def _render_pdf_with_playwright(html_str: str) -> bytes:
    """
    Render PDF using Playwright + headless Chromium.
    NOTE: On Railway you MUST have system libs installed (libglib2.0 etc) via Dockerfile.
    """
    _prepare_windows_event_loop()

    try:
        from playwright.sync_api import sync_playwright  # type: ignore
    except ImportError as e:
        raise RuntimeError("Playwright is not installed. Add `playwright` to requirements.txt.") from e

    logger.info("[PDF] Generating PDF with Playwright/Chromium")

    try:
        with sync_playwright() as p:
            browser = p.chromium.launch(
                headless=True,
                args=["--no-sandbox", "--disable-dev-shm-usage"],
            )
            page = browser.new_page()

            encoded_html = quote(html_str)
            data_url = f"data:text/html;charset=utf-8,{encoded_html}"

            page.goto(data_url, wait_until="domcontentloaded", timeout=30_000)

            pdf_bytes = page.pdf(
                format="A4",
                print_background=True,
                margin={"top": "12mm", "bottom": "12mm", "left": "12mm", "right": "12mm"},
            )

            browser.close()
            return pdf_bytes

    except Exception as e:
        logger.exception("[PDF] Playwright/Chromium PDF generation failed")
        raise RuntimeError(f"Playwright/Chromium PDF generation failed: {e}") from e


def render_cv_pdf_bytes(cv: CV, template_name: str) -> bytes:
    html_str = render_cv_html(cv, template_name=template_name)
    return _render_pdf_with_playwright(html_str)


# ============================================================
# DOCX: CV
# ============================================================
def render_cv_docx_bytes(cv: CV) -> bytes:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, Inches

    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title_p.add_run(cv.full_name or "Curriculum Vitae")
    r.bold = True
    r.font.size = Pt(16)

    if cv.title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(cv.title)

    contact_bits = []
    if getattr(cv, "email", None):
        contact_bits.append(cv.email)
    if getattr(cv, "phone", None):
        contact_bits.append(cv.phone)
    if getattr(cv, "full_address", None):
        contact_bits.append(cv.full_address)
    elif getattr(cv, "location", None):
        contact_bits.append(cv.location)

    if contact_bits:
        contact_p = doc.add_paragraph(" | ".join(contact_bits))
        contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_p.paragraph_format.space_after = Pt(12)

    if getattr(cv, "summary", None):
        doc.add_heading("Profile", level=2)
        doc.add_paragraph(cv.summary)

    if getattr(cv, "skills", None) and cv.skills:
        doc.add_heading("Skills", level=2)
        doc.add_paragraph(", ".join(cv.skills))

    if getattr(cv, "experiences", None) and cv.experiences:
        doc.add_heading("Experience", level=2)
        for exp in cv.experiences:
            header_parts = []
            if exp.job_title:
                header_parts.append(exp.job_title)
            if exp.company:
                header_parts.append(exp.company)
            header = " – ".join(header_parts)

            if header:
                h = doc.add_paragraph()
                h.add_run(header).bold = True

            meta_bits = []
            if exp.location:
                meta_bits.append(exp.location)

            date_bits = []
            if exp.start_date:
                date_bits.append(exp.start_date)
            if exp.end_date:
                date_bits.append(exp.end_date)
            if date_bits:
                meta_bits.append(" – ".join(date_bits))

            if meta_bits:
                doc.add_paragraph(" | ".join(meta_bits))

            if exp.description:
                for line in exp.description.splitlines():
                    line = line.strip("• ").strip()
                    if line:
                        doc.add_paragraph(line, style="List Bullet")

    if getattr(cv, "education", None) and cv.education:
        doc.add_heading("Education", level=2)
        for edu in cv.education:
            p = doc.add_paragraph()
            if edu.degree:
                p.add_run(edu.degree).bold = True

            meta_bits = []
            if edu.institution:
                meta_bits.append(edu.institution)
            if edu.location:
                meta_bits.append(edu.location)

            date_bits = []
            if edu.start_date:
                date_bits.append(edu.start_date)
            if edu.end_date:
                date_bits.append(edu.end_date)
            if date_bits:
                meta_bits.append(" – ".join(date_bits))

            if meta_bits:
                doc.add_paragraph(" | ".join(meta_bits))

    if getattr(cv, "references", None):
        text = cv.references or ""
        if text.strip():
            doc.add_heading("References", level=2)
            for line in text.splitlines():
                if line.strip():
                    doc.add_paragraph(line.strip())

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()


# ============================================================
# PDF: Cover letter (Playwright-only)
# ============================================================
def render_cover_letter_pdf_bytes(
    full_name: str,
    letter_body: str,
    location: str = "",
    email: str = "",
    phone: str = "",
    employer_name: str = "",
    employer_company: str = "",
    employer_location: str = "",
    greeting_name: str = "Hiring Manager",
    template_name: str = "cover_letter_basic.html",
) -> bytes:
    today_str = date.today().strftime("%d %B %Y")

    cleaned_lines = []
    for line in (letter_body or "").splitlines():
        s = line.strip()
        if not s:
            cleaned_lines.append("")
            continue
        lower = s.lower()
        if (
            (full_name and full_name.lower() in lower)
            or (email and email.lower() in lower)
            or (phone and phone in s)
            or (location and location.lower() in lower)
            or (today_str in s)
        ):
            continue
        cleaned_lines.append(line)

    cleaned_body = "\n".join(cleaned_lines).strip()

    template = env.get_template(template_name)
    html_str = template.render(
        full_name=full_name,
        location=location,
        email=email,
        phone=phone,
        employer_name=employer_name,
        employer_company=employer_company,
        employer_location=employer_location,
        greeting_name=greeting_name,
        today=today_str,
        letter_body=cleaned_body,
    )

    return _render_pdf_with_playwright(html_str)


# ============================================================
# DOCX: Cover letter
# ============================================================
def render_cover_letter_docx_bytes(
    full_name: str,
    letter_body: str,
    location: str = "",
    email: str = "",
    phone: str = "",
    employer_name: str = "",
    employer_company: str = "",
    employer_location: str = "",
    greeting_name: str = "Hiring Manager",
) -> bytes:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, Inches
    import re as _re

    today_str = date.today().strftime("%d %B %Y")

    def _is_headerish(s: str) -> bool:
        text = s.strip()
        if not text:
            return False
        words = text.replace(",", " ").split()
        return 1 < len(words) <= 6 and not text.endswith(".")

    location_tokens = []
    if location:
        location_tokens = [t.strip().lower() for t in _re.split(r"[,/]", location) if t.strip()]

    cleaned_lines = []
    for line in (letter_body or "").splitlines():
        s = line.strip()
        if not s:
            cleaned_lines.append("")
            continue
        lower = s.lower()

        location_hit = any(tok in lower for tok in location_tokens) if location_tokens else False

        headerish_employer = _is_headerish(s) and (
            (employer_company and employer_company.lower() in lower)
            or (employer_location and employer_location.lower() in lower)
            or (employer_name and employer_name.lower() in lower and "dear" not in lower)
        )

        if (
            (full_name and full_name.lower() in lower)
            or (email and email.lower() in lower)
            or (phone and phone in s)
            or (location_hit and _is_headerish(s))
            or headerish_employer
            or (today_str in s)
        ):
            continue

        cleaned_lines.append(line)

    cleaned_body = "\n".join(cleaned_lines).strip()

    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"

    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(6)

    header.add_run(full_name + "\n").bold = True
    if location:
        header.add_run(location + "\n")
    if email:
        header.add_run(email + "\n")
    if phone:
        header.add_run(phone + "\n")

    date_p = doc.add_paragraph(today_str)
    date_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    date_p.paragraph_format.space_after = Pt(6)

    if employer_name or employer_company or employer_location:
        emp = doc.add_paragraph()
        emp.paragraph_format.space_after = Pt(6)
        if employer_name:
            emp.add_run(employer_name + "\n")
        if employer_company:
            emp.add_run(employer_company + "\n")
        if employer_location:
            emp.add_run(employer_location + "\n")

    greet = doc.add_paragraph(f"Dear {greeting_name},")
    greet.paragraph_format.space_after = Pt(12)

    for paragraph in cleaned_body.split("\n\n"):
        cleaned = paragraph.strip()
        if cleaned:
            p = doc.add_paragraph(cleaned)
            p.paragraph_format.space_after = Pt(6)

    closing = doc.add_paragraph("Kind regards,")
    closing.paragraph_format.space_after = Pt(0)
    name_p = doc.add_paragraph(full_name)
    name_p.paragraph_format.space_before = Pt(0)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
